import asyncio
import json
import os
from datetime import datetime
from pathlib import Path
from typing import List

from data_model import ExportSet
from docx import Document


def _export_json(question_sets: List[ExportSet]):
    now = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = Path(f"{now}.json")
    with open(filename, "w") as f:
        json.dump([question_set.dict() for question_set in question_sets], f)
    return filename


def _export_txt(question_sets: List[ExportSet]):
    now = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = Path(f"{now}.txt")
    with open(filename, "w") as f:
        for question_set in question_sets:
            f.write(f"{question_set.context}\n\n")
            for qa_pair in question_set.question_pairs:
                f.write(f"{qa_pair.question}\n\n")
                for option in qa_pair.options:
                    if option.is_answer:
                        f.write(f"* {option.text}\n")
                    else:
                        f.write(f"- {option.text}\n")
                f.write("\n")
            f.write("\n\n")
        f.write("\n\n\n")
    return filename


def _export_docx(question_sets: List[ExportSet]):
    now = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = Path(f"{now}.docx")

    document = Document()
    for question_set in question_sets:
        document.add_paragraph(f"{question_set.context}")
        document.add_paragraph("")
        for qa_pair in question_set.question_pairs:
            document.add_paragraph(f"{qa_pair.question}")
            document.add_paragraph("")
            for option in qa_pair.options:
                if option.is_answer:
                    document.add_paragraph(f"+ {option.text}")
                else:
                    document.add_paragraph(f"- {option.text}")
            document.add_paragraph("")
    document.save(filename)

    return filename


def export_file(question_sets: List[ExportSet], format: str):
    try:
        return globals()[f"_export_{format.lower()}"](question_sets)
    except KeyError:
        raise ValueError(f"Unsupported format: {format.lower()}")


async def delete_later(file_path, wait=120):
    await asyncio.sleep(wait)
    os.remove(file_path)
